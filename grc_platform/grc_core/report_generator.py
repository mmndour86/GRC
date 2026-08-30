"""
Génération automatique du rapport Word d'évaluation de maturité cybersécurité
PSSI-ES, sur le modèle des rapports ANAQ-Sup / COUD fournis en exemple.

Usage: generer_rapport(conn, controle_id, chemin_sortie) -> chemin du .docx généré
"""
import io
import json
import os
import tempfile
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from grc_core.scoring import compute_scores, top_chapitres_faibles, top_objectifs_faibles, NIVEAU_LABELS
from grc_core.risk_scoring import get_registre_risques
from grc_core.db import DATA_DIR, nom_chapitre_affichage

NAVY = RGBColor(0x1F, 0x38, 0x64)
GOLD = RGBColor(0xB0, 0x8A, 0x2E)
GREY = RGBColor(0x44, 0x44, 0x44)

with open(os.path.join(DATA_DIR, "recommandations.json"), encoding="utf-8") as f:
    RECOMMANDATIONS = {int(k): v for k, v in json.load(f).items()}

# Objectifs dont les recommandations proviennent telles quelles du classeur Excel
# transmis par la DCSSI (les autres sont des propositions générées par la plateforme).
SOURCE_OFFICIELLE = {1}


def _libelle_recs(num, recs):
    suffixe = " (recommandation officielle DCSSI)" if num in SOURCE_OFFICIELLE else " (proposition à valider)"
    return " ; ".join(recs) + suffixe

with open(os.path.join(DATA_DIR, "mapping_referentiels.json"), encoding="utf-8") as f:
    MAPPING = {m["chapitre_num"]: m for m in json.load(f)}


# ---------------------------------------------------------------------------
# Utilitaires de mise en forme
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc, text, level=1, color=NAVY, size=16):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def _add_para(doc, text, size=11, bold=False, italic=False, color=None, space_after=8, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def _add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[i], "1F3864")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def _add_toc_field(doc):
    """Insère un champ TOC (table des matières) réel, à mettre à jour dans Word
    (clic droit -> Mettre à jour les champs) ou automatiquement à l'ouverture."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Clic droit ici puis « Mettre à jour les champs » pour générer le sommaire."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def _bar_chart_chapitres(scores, tmpdir):
    chap = scores["chapitres"]
    nums = sorted(chap.keys())
    labels = [f"Ch.{n}" for n in nums]
    values = [chap[n]["score"] or 0 for n in nums]
    colors = ["#c0392b" if v < 1.5 else "#e67e22" if v < 3 else "#27ae60" for v in values]

    fig, ax = plt.subplots(figsize=(7.2, 3.4), dpi=150)
    ax.bar(labels, values, color=colors)
    ax.set_ylim(0, 5.9)
    ax.set_yticks(range(0, 6))
    ax.set_ylabel("Maturité moyenne / 5")
    ax.set_title("Niveau de maturité par chapitre PSSI-ES", pad=12)
    ax.axhline(y=sum(values) / len(values) if values else 0, color="grey", linestyle="--", linewidth=0.8)
    for i, v in enumerate(values):
        ax.text(i, v + 0.15, f"{v:.2f}", ha="center", fontsize=8)
    plt.xticks(rotation=45, ha="right", fontsize=8)
    plt.tight_layout()
    path = os.path.join(tmpdir, "chap_scores.png")
    fig.savefig(path)
    plt.close(fig)
    return path


def _pie_repartition(scores, tmpdir):
    rep = scores["global"]["repartition_niveaux"]
    labels = [NIVEAU_LABELS[k] for k in rep]
    values = list(rep.values())
    if not values:
        return None
    fig, ax = plt.subplots(figsize=(4.5, 4.5), dpi=150)
    palette = {"Aucun": "#7f1d1d", "Initial": "#c0392b", "Géré": "#e67e22",
               "Défini": "#f1c40f", "Quantifié": "#2ecc71", "Optimisé": "#27ae60"}
    colors = [palette.get(l, "#95a5a6") for l in labels]
    ax.pie(values, labels=[f"{l} ({v})" for l, v in zip(labels, values)], colors=colors,
           autopct="%1.0f%%", startangle=90, textprops={"fontsize": 8})
    ax.set_title("Répartition des règles par niveau de maturité")
    plt.tight_layout()
    path = os.path.join(tmpdir, "repartition.png")
    fig.savefig(path)
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# Génération de texte analytique (imite le style rédactionnel des rapports fournis)
# ---------------------------------------------------------------------------

def _phrase_constat_central(scores, nom_entite):
    chap = scores["chapitres"]
    faibles = top_chapitres_faibles(scores, n=3)
    forts = sorted(
        [(n, c) for n, c in chap.items() if c["score"] is not None],
        key=lambda x: -x[1]["score"],
    )[:2]
    bullets = []
    for num, c in faibles:
        bullets.append(
            f"Le chapitre {num} « {nom_chapitre_affichage(num, c['nom'])} » affiche une maturité de "
            f"{c['score']:.2f}/5 ({c['n_repondues']} règle(s) évaluée(s) sur {c['n_regles']}, "
            f"{c['n_na']} non applicable(s))."
        )
    if forts:
        noms = " et ".join(f"« {nom_chapitre_affichage(n, c['nom'])} » ({c['score']:.2f}/5)" for n, c in forts)
        bullets.append(f"À l'inverse, {noms} constituent les points d'appui les plus solides.")
    return bullets


def _priorites_action(scores, n=4):
    """Génère les priorités d'action recommandées à partir des chapitres/objectifs
    les plus faibles, en s'appuyant sur le catalogue de recommandations quand il
    est renseigné, et sur un libellé générique dérivé de l'objectif sinon."""
    priorites = []
    objectifs_faibles = top_objectifs_faibles(scores, n=n)
    for num, o in objectifs_faibles:
        recs = RECOMMANDATIONS.get(num)
        if recs:
            priorites.append(f"Objectif {num} ({o['score']:.2f}/5) : " + _libelle_recs(num, recs))
        else:
            texte_court = o["texte"].strip().rstrip(".")
            if len(texte_court) > 160:
                texte_court = texte_court[:157] + "..."
            priorites.append(
                f"Objectif {num} ({o['score']:.2f}/5) : prioriser la mise en œuvre de « {texte_court} »."
            )
    return priorites


# ---------------------------------------------------------------------------
# Fonction principale
# ---------------------------------------------------------------------------

def generer_rapport(conn, controle_id, chemin_sortie, logo_path=None):
    controle = conn.execute(
        "SELECT c.*, e.nom AS entite_nom, e.secteur FROM controles c "
        "JOIN entites e ON e.id = c.entite_id WHERE c.id = ?", (controle_id,)
    ).fetchone()
    if controle is None:
        raise ValueError("Contrôle introuvable")

    nom_entite = controle["entite_nom"]
    scores = compute_scores(conn, controle_id)
    g = scores["global"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Page de garde ---
    for line, size, bold in [
        ("PRÉSIDENCE DE LA RÉPUBLIQUE DU SÉNÉGAL", 12, True),
        ("SECRÉTARIAT GÉNÉRAL", 11, False),
        ("DIRECTION GÉNÉRALE DU CHIFFRE ET DE LA SÉCURITÉ DES SYSTÈMES D'INFORMATION (DCSSI)", 11, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RAPPORT D'ÉVALUATION")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"DU NIVEAU DE MATURITÉ CYBERSÉCURITÉ DE {nom_entite.upper()}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = GOLD

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Évaluation de la mise en œuvre des règles de la Politique de Sécurité des Systèmes "
        "d'Information de l'État du Sénégal (PSSI-ES)"
    )
    r.italic = True
    r.font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Rapport généré automatiquement le {date.today().strftime('%d/%m/%Y')} "
                  f"par la plateforme GRC — contrôle n°{controle_id}")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

    doc.add_page_break()

    if controle["source_donnees"] == "reconstitution":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        run = p.add_run(
            "⚠ DONNÉES RECONSTITUÉES — Ce contrôle recharge les résultats d'une évaluation DCSSI déjà "
            "publiée pour cette entité, à partir des scores moyens par objectif figurant dans le rapport "
            "d'origine (celui-ci ne conservant pas les réponses détaillées règle par règle). Les scores "
            "par objectif, par chapitre et le score global reproduisent fidèlement le rapport d'origine ; "
            "le niveau affiché pour chaque règle individuelle est une approximation calculée par la "
            "plateforme et ne doit pas être considéré comme la réponse effectivement déclarée à l'époque."
        )
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xB0, 0x2E, 0x2E)
        doc.add_page_break()

    # --- Sommaire ---
    _add_heading(doc, "Sommaire", level=1)
    _add_toc_field(doc)
    doc.add_page_break()

    tmpdir = tempfile.mkdtemp()

    # --- Résumé exécutif ---
    _add_heading(doc, "Résumé exécutif", level=1)
    _add_para(
        doc,
        f"L'évaluation de la maturité cybersécurité de {nom_entite}, conduite sur l'intégralité des "
        f"{g['n_regles']} règles de la PSSI-ES, fait ressortir un score global de {g['score']:.2f}/5 "
        f"(soit un taux de maturité de {g['taux_conformite']:.1f} %), correspondant au niveau CMMI "
        f"« {g['niveau_label']} »." if g["score"] is not None else
        f"L'évaluation de {nom_entite} est en cours : {g['n_repondues']} règle(s) sur {g['n_regles']} "
        f"ont été renseignées à ce stade."
    )

    _add_heading(doc, "Constat central", level=2)
    for phrase in _phrase_constat_central(scores, nom_entite):
        _add_bullet(doc, phrase)

    _add_heading(doc, "Chiffres clés à retenir", level=2)
    _add_bullet(doc, f"Score global de maturité : {g['score']:.2f}/5 ({g['taux_conformite']:.1f} %)." if g["score"] is not None else "Score global : évaluation incomplète.")
    _add_bullet(doc, f"{g['n_repondues']} règle(s) évaluée(s), {g['n_na']} déclarée(s) non applicable(s), {g['n_non_renseignees']} non renseignée(s) sur {g['n_regles']}.")
    for niveau, label in NIVEAU_LABELS.items():
        c = g["repartition_niveaux"].get(niveau, 0)
        if c:
            pct = round(c / g["n_repondues"] * 100, 1) if g["n_repondues"] else 0
            _add_bullet(doc, f"{c} règle(s) au niveau « {label} » ({pct} % des règles évaluées).")
    chap_zero = [c["nom"] for c in scores["chapitres"].values() if c["score"] == 0]
    if chap_zero:
        _add_bullet(doc, f"Chapitre(s) au score de 0/5 : {', '.join(chap_zero)}.")

    _add_heading(doc, "Priorités d'action recommandées à la direction", level=2)
    _add_para(
        doc,
        "Sauf mention contraire (« recommandation officielle DCSSI »), les recommandations ci-dessous "
        "sont des propositions générées par la plateforme à partir du libellé des objectifs PSSI-ES et "
        "des bonnes pratiques usuelles ; elles sont à valider et à adapter par vos équipes conformité "
        "avant diffusion officielle.",
        size=9, italic=True, color=GREY, space_after=10,
    )
    for i, priorite in enumerate(_priorites_action(scores), start=1):
        _add_bullet(doc, f"{i}. {priorite}")

    doc.add_page_break()

    # --- 1. Contexte ---
    _add_heading(doc, "1. Contexte et objectifs de la mission", level=1)
    _add_para(
        doc,
        "Dans le cadre de ses missions de régulation et de contrôle de la sécurité des systèmes "
        "d'information de l'État du Sénégal, la Direction générale du Chiffre et de la Sécurité des "
        f"Systèmes d'Information (DCSSI) a évalué le niveau de maturité cybersécurité de {nom_entite} "
        "au regard de la Politique de Sécurité des Systèmes d'Information de l'État du Sénégal (PSSI-ES)."
    )
    _add_para(
        doc,
        "Cette évaluation s'appuie sur le questionnaire d'auto-évaluation renseigné dans la plateforme "
        "GRC, qui documente pour chacune des règles applicables le niveau de maturité atteint, "
        "les justificatifs de non-applicabilité et les commentaires associés."
    )
    _add_para(doc, "L'objectif du présent rapport est de :")
    _add_bullet(doc, f"dresser un état des lieux exhaustif et objectif du niveau de maturité cybersécurité de {nom_entite} au regard des {g['n_regles']} règles de la PSSI-ES ;")
    _add_bullet(doc, "mettre en évidence les points forts et les axes de progrès, chapitre par chapitre et objectif par objectif ;")
    _add_bullet(doc, "formuler des recommandations priorisées, assorties d'un plan d'action opérationnel.")

    # --- 2. Méthodologie ---
    _add_heading(doc, "2. Méthodologie d'évaluation", level=1)
    _add_para(
        doc,
        "L'évaluation repose sur une grille de maturité inspirée du modèle CMMI (Capability Maturity "
        "Model Integration), adaptée par la DCSSI aux exigences de la PSSI-ES. Six niveaux sont utilisés : "
        "Aucun (0), Initial (1), Géré (2), Défini (3), Quantifié (4) et Optimisé (5)."
    )
    _add_para(
        doc,
        "Le taux de conformité par objectif et par chapitre est calculé automatiquement par la "
        "plateforme GRC sur la base de la moyenne des niveaux de maturité des règles applicables et "
        "renseignées, rapportée au score maximal de 5. Les règles déclarées non applicables sont "
        "exclues du calcul et font l'objet d'un suivi de justification séparé."
    )

    # --- 3. Synthèse des résultats globaux ---
    _add_heading(doc, "3. Synthèse des résultats globaux", level=1)
    _add_para(
        doc,
        f"L'évaluation de {nom_entite} porte sur l'intégralité des {g['n_regles']} règles de la PSSI-ES, "
        "réparties en 30 objectifs de sécurité regroupés au sein de 11 chapitres."
    )

    _add_heading(doc, "3.1 Répartition des règles par niveau de maturité", level=2)
    pie_path = _pie_repartition(scores, tmpdir)
    if pie_path:
        doc.add_picture(pie_path, width=Cm(10))

    _add_heading(doc, "3.2 Niveau de maturité par chapitre", level=2)
    chart_path = _bar_chart_chapitres(scores, tmpdir)
    doc.add_picture(chart_path, width=Cm(16))
    rows = []
    for num in sorted(scores["chapitres"].keys()):
        c = scores["chapitres"][num]
        score_txt = f"{c['score']:.2f}/5" if c["score"] is not None else "Non évalué"
        rows.append([num, nom_chapitre_affichage(num, c["nom"]), score_txt, c["n_repondues"], c["n_na"]])
    _add_table(doc, ["Chapitre n°", "Intitulé", "Maturité moyenne", "Règles évaluées", "Non applicables"], rows,
               col_widths=[2, 8, 3, 2.5, 2.5])

    _add_heading(doc, "3.3 Niveau de maturité par objectif de sécurité", level=2)
    rows = []
    for num in sorted(scores["objectifs"].keys()):
        o = scores["objectifs"][num]
        score_txt = f"{o['score']:.2f}/5" if o["score"] is not None else "Non évalué"
        texte_court = o["texte"][:100] + ("..." if len(o["texte"]) > 100 else "")
        rows.append([num, o["chapitre_num"], texte_court, score_txt])
    _add_table(doc, ["Objectif n°", "Chapitre", "Libellé", "Maturité moyenne"], rows,
               col_widths=[2, 2, 11, 3])

    doc.add_page_break()

    # --- 4. Constats détaillés par chapitre ---
    _add_heading(doc, "4. Constats détaillés par chapitre", level=1)
    _add_para(
        doc,
        "Cette section reprend, pour chacun des 11 chapitres de la PSSI-ES, le niveau de maturité "
        "moyen constaté, la couverture de l'évaluation et la correspondance avec les référentiels "
        "ISO 27001, NIST CSF et DORA."
    )
    for num in sorted(scores["chapitres"].keys()):
        c = scores["chapitres"][num]
        score_txt = f"{c['score']:.2f}/5" if c["score"] is not None else "non évalué"
        _add_heading(doc, f"{num}. {nom_chapitre_affichage(num, c['nom'])} — maturité moyenne : {score_txt}", level=2)
        _add_para(
            doc,
            f"{c['n_repondues']} règle(s) évaluée(s) sur {c['n_regles']}, dont {c['n_na']} déclarée(s) "
            "non applicable(s)."
        )
        m = MAPPING.get(num)
        if m:
            rows = [
                ["ISO/IEC 27001", m["iso27001"]],
                ["NIST CSF", m["nist_csf"]],
                ["DORA", m["dora"]],
            ]
            _add_table(doc, ["Référentiel", "Correspondance"], rows, col_widths=[3, 13])
        objectifs_du_chapitre = [
            (n, o) for n, o in scores["objectifs"].items() if o["chapitre_num"] == num
        ]
        for onum, o in sorted(objectifs_du_chapitre):
            recs = RECOMMANDATIONS.get(onum)
            if recs and o["score"] is not None and o["score"] < 3:
                suffixe = "officielle DCSSI" if onum in SOURCE_OFFICIELLE else "proposée, à valider"
                _add_para(doc, f"Objectif {onum} — recommandations ({suffixe}) :", bold=True, size=10)
                for rec in recs:
                    _add_bullet(doc, rec, size=10)

    doc.add_page_break()

    # --- 5. Plan d'action ---
    _add_heading(doc, "5. Plan d'action de mise en œuvre", level=1)
    _add_para(
        doc,
        "Le plan d'action ci-dessous priorise les objectifs affichant le score de maturité le plus "
        "faible. Les colonnes responsable, échéance et budget sont à compléter dans la plateforme lors "
        "du suivi opérationnel."
    )
    _add_para(
        doc,
        "Colonne « Source » : DCSSI = recommandation officielle reprise du référentiel transmis ; "
        "Proposée = action générée par la plateforme à partir du libellé de l'objectif, à valider.",
        size=9, italic=True, color=GREY, space_after=8,
    )
    rows = []
    for onum, o in top_objectifs_faibles(scores, n=10):
        chap_nom = nom_chapitre_affichage(o["chapitre_num"], scores["chapitres"][o["chapitre_num"]]["nom"])
        recs = RECOMMANDATIONS.get(onum)
        action = " ; ".join(recs) if recs else f"Formaliser et mettre en œuvre : {o['texte'][:120]}"
        score_txt = f"{o['score']:.2f}/5" if o["score"] is not None else "N/A"
        source = "DCSSI" if onum in SOURCE_OFFICIELLE else "Proposée"
        rows.append([chap_nom, f"Objectif {onum}", action[:150], source, score_txt, "", ""])
    _add_table(
        doc,
        ["Chapitre", "Objectif", "Action à réaliser", "Source", "Maturité actuelle", "Responsable", "Échéance"],
        rows, col_widths=[3, 1.7, 6.5, 1.6, 1.7, 2, 2],
    )

    # --- Annexe risques (si des risques ISO 27005 existent pour cette entité) ---
    risques = get_registre_risques(conn, controle["entite_id"])
    if risques:
        doc.add_page_break()
        _add_heading(doc, "6. Annexe — Synthèse des risques (ISO/IEC 27005)", level=1)
        _add_para(
            doc,
            "Cette annexe complète l'évaluation de conformité par un extrait du registre des risques "
            "de sécurité de l'information de l'entité, apprécié selon la méthodologie ISO/IEC 27005."
        )
        risques_tries = sorted(risques, key=lambda r: -(r["niveau_risque_residuel"] or 0))
        rows = []
        for r in risques_tries[:15]:
            rows.append([
                r["code"], (r["description_scenario"] or "")[:90],
                r["niveau_qualitatif_brut"] or "-", r["niveau_qualitatif_residuel"] or "-",
                "Oui" if r.get("depasse_appetence") else "Non",
            ])
        _add_table(doc, ["ID", "Scénario de risque", "Niveau brut", "Niveau résiduel", "Dépasse l'appétence"],
                   rows, col_widths=[1.5, 10, 2.5, 2.5, 2.5])

    # --- Conclusion ---
    doc.add_page_break()
    _add_heading(doc, "7. Conclusion", level=1)
    if g["score"] is not None:
        _add_para(
            doc,
            f"Le score global de {g['score']:.2f}/5 place {nom_entite} au niveau « {g['niveau_label']} » "
            "de l'échelle de maturité CMMI. La mise en œuvre du plan d'action ci-dessus, en priorisant "
            "les chapitres et objectifs les plus faibles, permettra de faire progresser durablement le "
            "niveau de maturité cybersécurité de l'entité. Un nouveau contrôle est recommandé pour "
            "mesurer les progrès réalisés."
        )
    else:
        _add_para(doc, "L'évaluation étant partielle, ce rapport devra être régénéré une fois le questionnaire complété.")

    _add_para(
        doc,
        "Rapport généré automatiquement par la plateforme GRC à partir des données saisies dans le module "
        "de conformité PSSI-ES et le registre des risques ISO/IEC 27005.",
        size=9, italic=True, color=GREY, space_after=0,
    )

    doc.save(chemin_sortie)
    return chemin_sortie
