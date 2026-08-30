"""Tests du générateur de rapport Word (grc_core.report_generator)."""
import os

from docx import Document

from grc_core.report_generator import generer_rapport


def _premier_controle(conn, nom_entite):
    row = conn.execute(
        "SELECT c.id FROM controles c JOIN entites e ON e.id=c.entite_id WHERE e.nom=? LIMIT 1",
        (nom_entite,),
    ).fetchone()
    return row["id"] if row else None


def test_generer_rapport_produit_un_docx_valide(conn, tmp_path):
    controle_id = _premier_controle(conn, "Entité de démonstration")
    assert controle_id is not None

    out_path = tmp_path / "rapport_test.docx"
    generer_rapport(conn, controle_id, str(out_path))

    assert out_path.exists()
    assert out_path.stat().st_size > 10_000  # un .docx avec graphiques n'est jamais minuscule

    # Le fichier doit être un .docx exploitable par python-docx (pas juste des
    # octets présents sur disque).
    doc = Document(str(out_path))
    assert len(doc.paragraphs) > 0


def test_rapport_reconstitution_contient_le_bandeau_avertissement(conn, tmp_path):
    """Les entités ANAQ-Sup/COUD reconstituées doivent porter un avertissement
    explicite de provenance dans le rapport généré."""
    row = conn.execute(
        "SELECT c.id, e.nom FROM controles c JOIN entites e ON e.id=c.entite_id "
        "WHERE c.source_donnees='reconstitution' LIMIT 1"
    ).fetchone()
    if row is None:
        import pytest
        pytest.skip("Aucune entité reconstituée dans ce jeu de données de test.")

    out_path = tmp_path / "rapport_reconstitution.docx"
    generer_rapport(conn, row["id"], str(out_path))

    doc = Document(str(out_path))
    texte_complet = "\n".join(p.text for p in doc.paragraphs)
    assert "reconstit" in texte_complet.lower()


def test_rapport_contient_le_nom_de_lentite(conn, tmp_path):
    controle_id = _premier_controle(conn, "Entité de démonstration")
    out_path = tmp_path / "rapport_nom.docx"
    generer_rapport(conn, controle_id, str(out_path))

    doc = Document(str(out_path))
    texte_complet = "\n".join(p.text for p in doc.paragraphs)
    assert "Entité de démonstration" in texte_complet
